#!/usr/bin/env python3
"""Generate practice diary (дневник практики) for OrderFlow internship."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Margins: left 30mm, right 10mm, top/bottom 20mm
for s in doc.sections:
    s.top_margin    = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin   = Cm(3)
    s.right_margin  = Cm(1)

F = 'Times New Roman'

def p(text='', align=WD_ALIGN_PARAGRAPH.JUSTIFY, sz=14, bold=False,
      indent=True, sb=0, sa=4):
    par = doc.add_paragraph()
    par.alignment = align
    par.paragraph_format.space_before = Pt(sb)
    par.paragraph_format.space_after  = Pt(sa)
    if indent and align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        par.paragraph_format.first_line_indent = Cm(1.25)
    if text:
        r = par.add_run(text)
        r.font.name = F; r.font.size = Pt(sz); r.font.bold = bold
    return par

def center(text, sz=14, bold=False, sb=4, sa=4):
    return p(text, align=WD_ALIGN_PARAGRAPH.CENTER,
             sz=sz, bold=bold, indent=False, sb=sb, sa=sa)

def left(text, sz=14, bold=False, sb=2, sa=2):
    return p(text, align=WD_ALIGN_PARAGRAPH.LEFT,
             sz=sz, bold=bold, indent=False, sb=sb, sa=sa)

def pb(): doc.add_page_break()

def blank_line(text='_' * 42, sz=14):
    return left(text, sz=sz)

def add_line_with_label(label, value='', sz=14):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    par.paragraph_format.space_after = Pt(4)
    r1 = par.add_run(label)
    r1.font.name = F; r1.font.size = Pt(sz)
    r2 = par.add_run(value if value else ' ' + '_' * (50 - len(label)))
    r2.font.name = F; r2.font.size = Pt(sz)
    return par

# ═══════════════════════════════════════════════════════════════════════════
# TITLE / HEADER PAGE
# ═══════════════════════════════════════════════════════════════════════════
center('ЧОУ ВО «Казанский инновационный университет', sz=12, sb=0)
center('имени В.Г. Тимирясова (ИЭУП)»', sz=12)
center('Факультет менеджмента и инженерного бизнеса', sz=12, sb=6)
p('')
p('')
center('ДНЕВНИК ПРАКТИКИ', sz=16, bold=True, sb=20, sa=6)
p('')
left('Вид и тип практики: производственная (преддипломная)', sz=14, sb=6)
p('')
left('Фамилия ____________________________________', sz=14)
left('Имя, отчество _______________________________', sz=14)
left('Факультет менеджмента и инженерного бизнеса', sz=14)
left('Курс __4__    Группа ___1023___', sz=14)
left('Направление подготовки 09.03.03 Прикладная информатика', sz=14)
left('Профиль подготовки «Прикладная информатика в экономике»', sz=14)
left('Срок практики с 06.04.2026 г. по 19.05.2026 г.', sz=14)
p('')
p('')

# Инструктаж
center('Инструктаж по соблюдению правил внутреннего трудового', sz=13, bold=True, sb=10)
center('распорядка, требований охраны труда и пожарной безопасности', sz=13, bold=True)
p('')
left('Инструктаж провёл', sz=12)
left('Руководитель практики от профильной организации', sz=12)
p('')
left('_______________________________', sz=12)
left('_______________________________', sz=12)
left('(должность, ФИО)', sz=11)
left('М.П.', sz=12)
p('')
left('С требованиями охраны труда, техники безопасности, пожарной безопасности,', sz=12)
left('а также правилами внутреннего трудового распорядка ознакомлен(а).', sz=12)
p('')
left('________________________________________', sz=12)
left('(ФИО, подпись обучающегося)', sz=11)
p('')
left('Место прохождения практики: ООО «Парнас АйТи»', sz=14, sb=10)
left('Руководитель практики от профильной организации:', sz=14)
left('должность ___________________________________', sz=14)
left('ФИО _________________________________________', sz=14)
p('')
center('Казань, 2026', sz=12, sb=20)

pb()

# ═══════════════════════════════════════════════════════════════════════════
# MAIN DIARY TABLE
# ═══════════════════════════════════════════════════════════════════════════
center('Дневник прохождения практики', sz=14, bold=True, sb=6, sa=8)

# Diary entries: (n, date_str, work_description)
# Production calendar RF 2026:
# Holidays in Apr-May: May 1 (Fri - Labour Day), May 9 (Sat - Victory Day)
# Saturdays = working days per methodology
# Sundays = non-working
# Intervals ≤ 3 days

entries = [
    (1,
     '06 апр.',
     'Организационное собрание в университете. Получение бланка дневника практики, '
     'рабочего плана и индивидуального задания. Инструктаж по целям и порядку прохождения практики.'),

    (2,
     '07–09 апр.',
     'Прибытие в ООО «Парнас АйТи». Ознакомление с деятельностью организации, '
     'прохождение инструктажа по охране труда и правилам внутреннего распорядка. '
     'Изучение структуры отдела разработки, применяемых технологий и действующих проектов.'),

    (3,
     '10–11 апр.',
     'Анализ предметной области: изучение типичных процессов управления заказами '
     'на предприятии, выявление проблем ручного ведения заявок, определение круга '
     'участников бизнес-процесса.'),

    (4,
     '13–15 апр.',
     'Сравнительный анализ подходов к автоматизации управления заказами: инструменты '
     'общего назначения, готовые CRM/ERP-системы, разработка специализированной платформы. '
     'Обоснование выбора технологического стека (NestJS, PostgreSQL, Redis, Next.js).'),

    (5,
     '16–18 апр.',
     'Проектирование реляционной схемы базы данных: определение сущностей (заказы, '
     'пользователи, клиенты, позиции заказа, история изменений, уведомления), '
     'типов данных и связей между таблицами.'),

    (6,
     '20–22 апр.',
     'Разработка модуля аутентификации: регистрация и вход пользователей, '
     'JWT-аутентификация с access- и refresh-токенами, bcrypt-хеширование паролей, '
     'ролевое управление доступом (admin, manager, executor, client).'),

    (7,
     '23–25 апр.',
     'Разработка модуля управления заказами: CRUD-операции, статусная машина '
     'с контролем допустимых переходов между десятью состояниями, '
     'журнал истории изменений, автоматическая генерация номера заказа.'),

    (8,
     '27–29 апр.',
     'Разработка алгоритмов автоматического распределения заказов по исполнителям: '
     'Round Robin, минимальная загрузка (Min Load), приоритетное назначение. '
     'Интеграция алгоритмов в сервис обработки заказов.'),

    (9,
     '30 апр. – 2 мая',
     'Разработка модуля уведомлений: отправка Email-уведомлений через SMTP, '
     'системные in-app уведомления с хранением в базе данных. '
     'Разработка аналитического модуля: расчёт KPI-показателей, выручки, '
     'просроченных заказов и эффективности исполнителей.\n'
     '(1 мая — нерабочий праздничный день)'),

    (10,
     '04–06 мая',
     'Разработка административного модуля: управление учётными записями '
     'пользователей, назначение и изменение ролей, деактивация аккаунтов, '
     'системные настройки (алгоритм распределения, шаблоны уведомлений).'),

    (11,
     '07–08 мая',
     'Начало разработки фронтенда на Next.js 14: настройка проекта, '
     'страницы аутентификации (вход, регистрация), главный дашборд '
     'со статистическими карточками KPI.\n'
     '(9 мая — нерабочий праздничный день)'),

    (12,
     '11–13 мая',
     'Разработка страниц управления заказами: список с фильтрацией и поиском, '
     'детальная карточка заказа со сменой статуса и историей изменений. '
     'Разработка CRM-компонента: список клиентов, карточка клиента.'),

    (13,
     '14–16 мая',
     'Разработка аналитического дашборда: интерактивные графики динамики заказов, '
     'круговая диаграмма статусов, рейтинг исполнителей. '
     'Разработка панели администратора: управление пользователями и настройками.'),

    (14,
     '18 мая',
     'Настройка Docker Compose (сервисы: PostgreSQL, Redis, backend, frontend) '
     'и CI/CD-пайплайна на GitHub Actions. Комплексное тестирование платформы, '
     'устранение выявленных замечаний. Оформление отчёта по практике.'),

    (15,
     '19 мая',
     'Завершение оформления и сдача отчёта по практике. '
     'Защита отчёта по результатам преддипломной практики в университете.'),
]

# Build table
tbl = doc.add_table(rows=1 + len(entries), cols=3)
tbl.style = 'Table Grid'

# Header
hrow = tbl.rows[0]
col_labels = ['№ п/п', 'Дата', 'Рабочие записи']
widths     = [Cm(1.4), Cm(4.2), Cm(12.0)]
for i, (label, w) in enumerate(zip(col_labels, widths)):
    cell = hrow.cells[i]
    cell.width = w
    cell.text = ''
    rn = cell.paragraphs[0].add_run(label)
    rn.font.name = F; rn.font.size = Pt(12); rn.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Data rows
for row_idx, (num, date, work) in enumerate(entries):
    row = tbl.rows[row_idx + 1]
    row.cells[0].width = widths[0]
    row.cells[1].width = widths[1]
    row.cells[2].width = widths[2]

    # №
    row.cells[0].text = ''
    r0 = row.cells[0].paragraphs[0].add_run(str(num))
    r0.font.name = F; r0.font.size = Pt(12)
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    row.cells[1].text = ''
    r1 = row.cells[1].paragraphs[0].add_run(date)
    r1.font.name = F; r1.font.size = Pt(12); r1.font.bold = True
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Work notes — handle \n for multi-paragraph cells
    parts = work.split('\n')
    row.cells[2].text = ''
    first_par = row.cells[2].paragraphs[0]
    first_par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    first_par.paragraph_format.space_after = Pt(2)
    rn = first_par.add_run(parts[0])
    rn.font.name = F; rn.font.size = Pt(12)

    for extra in parts[1:]:
        new_par = row.cells[2].add_paragraph()
        new_par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        new_par.paragraph_format.space_after = Pt(2)
        rn2 = new_par.add_run(extra)
        rn2.font.name = F; rn2.font.size = Pt(11)
        rn2.font.italic = True

# Minimum row heights for readability
for row in tbl.rows[1:]:
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '600')
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

pb()

# ═══════════════════════════════════════════════════════════════════════════
# SUPERVISOR CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════
center('Заключение руководителя практики от профильной организации',
       sz=13, bold=True, sb=6, sa=8)

conclusion_rows = [
    'Соблюдение правил внутреннего трудового распорядка',
    'Соблюдение требований охраны труда и пожарной безопасности',
    'Соблюдение рабочего графика (плана) практики',
    'Владение способами и методами самоорганизации и самообразования\n'
    'в профессиональной деятельности',
    'Умение наладить сотрудничество в коллективе',
    'Умение применять теоретические знания в профессиональной деятельности',
    'Формирование практических умений и навыков, предусмотренных\n'
    'программой практики',
    'Выполнение индивидуального задания практики',
    'Добросовестность и активность при выполнении индивидуального задания',
]

ctbl = doc.add_table(rows=1 + len(conclusion_rows), cols=2)
ctbl.style = 'Table Grid'

# header
ch = ctbl.rows[0]
ch.cells[0].text = ''
rh0 = ch.cells[0].paragraphs[0].add_run('Показатели выполнения')
rh0.font.name = F; rh0.font.size = Pt(12); rh0.font.bold = True
ch.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
ch.cells[0].width = Cm(13)

ch.cells[1].text = ''
rh1 = ch.cells[1].paragraphs[0].add_run('Выполнено /\nНе выполнено')
rh1.font.name = F; rh1.font.size = Pt(12); rh1.font.bold = True
ch.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
ch.cells[1].width = Cm(4.6)

for i, label in enumerate(conclusion_rows):
    row = ctbl.rows[i + 1]
    row.cells[0].width = Cm(13)
    row.cells[1].width = Cm(4.6)
    parts = label.split('\n')
    row.cells[0].text = ''
    fp = row.cells[0].paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rn = fp.add_run(parts[0])
    rn.font.name = F; rn.font.size = Pt(12)
    for extra in parts[1:]:
        np2 = row.cells[0].add_paragraph()
        np2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rn2 = np2.add_run(extra)
        rn2.font.name = F; rn2.font.size = Pt(12)
    row.cells[1].text = ''

p('')
left('Руководитель практики от профильной организации', sz=13, sb=10)
left('___________________________________________________', sz=13)
left('(полное наименование организации/предприятия)', sz=11)
p('')
left('_________________  /  _______________________________', sz=13)
left('     (подпись)                        (ФИО)', sz=11)
left('М.П.', sz=13, sb=6)

# ─── Save ────────────────────────────────────────────────────────────────
out = '/home/user/Test/Дневник_практики_OrderFlow.docx'
doc.save(out)
print(f'Saved: {out}')
