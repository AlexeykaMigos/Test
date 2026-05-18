#!/usr/bin/env python3
"""Generate OrderFlow practice report as .docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Page margins ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(1.5)

# ─── Helper functions ────────────────────────────────────────────────────────
def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_paragraph(text='', align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  size=12, bold=False, italic=False,
                  space_before=0, space_after=6, first_indent=True):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if first_indent and align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        pf.first_line_indent = Cm(1.25)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p

def add_heading(text, level=1):
    sizes   = {1: 14, 2: 13, 3: 12}
    sz      = sizes.get(level, 12)
    before  = {1: 12, 2: 10, 3: 8}
    p = add_paragraph(text, align=WD_ALIGN_PARAGRAPH.CENTER,
                      size=sz, bold=True,
                      space_before=before.get(level, 8), space_after=8,
                      first_indent=False)
    return p

def add_bullet(text, size=12, indent=Cm(1.25)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent       = indent
    pf.first_line_indent = Cm(-0.5)
    pf.space_after       = Pt(3)
    run = p.add_run('– ' + text)
    set_font(run, size=size)
    return p

def add_numbered(text, num, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent       = Cm(1.25)
    pf.first_line_indent = Cm(-1.25)
    pf.space_after       = Pt(3)
    run = p.add_run(f'{num}. {text}')
    set_font(run, size=size)
    return p

def page_break():
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# TITLE PAGE  (blank – user will fill in)
# ═══════════════════════════════════════════════════════════════════════════
add_paragraph('', space_before=0, space_after=0)
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════════
add_heading('СОДЕРЖАНИЕ', level=1)

toc_items = [
    ('Введение', '3'),
    ('Раздел 1', ''),
    ('1. Характеристика предприятия и постановка задачи', '5'),
    ('1.1. Общая характеристика организации и подразделения', '5'),
    ('1.2. Техническое и программное обеспечение', '6'),
    ('1.3. Анализ существующих недостатков и проблем управления заказами', '7'),
    ('1.4. Формулировка цели и задач практики', '8'),
    ('Раздел 2', ''),
    ('2. Содержание мероприятий по разработке системы автоматизации заказов', '9'),
    ('2.1. Анализ и сравнительная оценка методов решения задач практики', '9'),
    ('2.2. Описание реализации решения', '11'),
    ('Заключение', '17'),
    ('Список использованных источников', '18'),
    ('Приложение А', '19'),
    ('Приложение Б', '21'),
    ('Приложение В', '22'),
]

for item, page in toc_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    indent = Cm(1.25) if item.startswith('1.') or item.startswith('2.') else Cm(0)
    if item.startswith('1.1') or item.startswith('1.2') or item.startswith('1.3') or item.startswith('1.4') or item.startswith('2.1') or item.startswith('2.2'):
        pf.left_indent = Cm(1.25)
    run = p.add_run(item)
    set_font(run, size=12)
    if page:
        run2 = p.add_run(f'..............{page}')
        set_font(run2, size=12)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════
add_heading('ВВЕДЕНИЕ', level=1)

add_paragraph(
    'ООО «Парнас АйТ» — российская IT-компания, специализирующаяся на разработке '
    'корпоративного программного обеспечения и цифровых платформ для автоматизации '
    'бизнес-процессов. Компания оказывает услуги по проектированию, разработке '
    'и внедрению информационных систем для предприятий различных отраслей: ритейла, '
    'логистики, производства и сферы услуг. В портфеле компании — решения '
    'для управления складом, CRM-системы, платформы для автоматизации документооборота '
    'и интеграционные сервисы.'
)

add_paragraph(
    'Основными направлениями деятельности компании являются:'
)

bullets_intro = [
    'Разработка корпоративного программного обеспечения на заказ',
    'Проектирование и внедрение SaaS-платформ для бизнеса',
    'Интеграция с внешними сервисами: CRM, платёжными системами, логистическими операторами',
    'Техническое сопровождение и поддержка разработанных решений',
    'Консалтинг в области цифровой трансформации бизнес-процессов',
]
for b in bullets_intro:
    add_bullet(b)

add_paragraph(
    'Практика проходила в отделе backend-разработки компании под руководством '
    'ведущего разработчика. Отдел отвечает за проектирование серверной архитектуры, '
    'разработку REST API, работу с базами данных и интеграцию внешних сервисов. '
    'Работа отдела охватывает полный цикл разработки и включает следующие направления:'
)

bullets_dept = [
    'Проектирование и реализация серверной бизнес-логики на Node.js / NestJS',
    'Проектирование реляционных схем баз данных PostgreSQL',
    'Разработка и документирование REST API (Swagger/OpenAPI)',
    'Реализация механизмов аутентификации и авторизации на базе JWT',
    'Настройка кэширования с помощью Redis',
    'Контейнеризация приложений с использованием Docker и Docker Compose',
    'Настройка CI/CD-пайплайнов на GitHub Actions',
]
for b in bullets_dept:
    add_bullet(b)

add_paragraph(
    'В рамках прохождения производственной практики была поставлена задача '
    'по проектированию и разработке веб-платформы OrderFlow — системы автоматизации '
    'полного цикла обработки заказов. Актуальность задачи обусловлена потребностью '
    'клиентов компании в централизованном инструменте управления заказами, который '
    'минимизирует ручные операции, ускоряет обработку, снижает количество ошибок '
    'и обеспечивает прозрачность бизнес-процессов через единый цифровой интерфейс.'
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 1
# ═══════════════════════════════════════════════════════════════════════════
add_heading('Раздел 1', level=1)
add_heading('1. ХАРАКТЕРИСТИКА ПРЕДПРИЯТИЯ И ПОСТАНОВКА ЗАДАЧИ', level=2)
add_heading('1.1. Общая характеристика организации и подразделения', level=3)

add_paragraph(
    'ООО «Парнас АйТ» является динамично развивающейся компанией в сфере '
    'информационных технологий, специализирующейся на разработке корпоративных '
    'информационных систем и веб-платформ. Компания основана в 2015 году '
    'и расположена в г. Санкт-Петербурге. За годы работы компания реализовала '
    'свыше 80 проектов для предприятий малого, среднего и крупного бизнеса '
    'в различных отраслях экономики.'
)

add_paragraph(
    'Ключевыми продуктами и направлениями деятельности компании являются:'
)

key_products = [
    'Разработка CRM-систем и платформ управления клиентами',
    'Создание систем автоматизации складской логистики и управления заказами',
    'Разработка интеграционных решений (API-шлюзы, ESB-шины)',
    'Внедрение корпоративных ERP-систем',
    'Разработка мобильных приложений и PWA-решений',
    'Аутстаффинг и предоставление команд разработчиков',
]
for b in key_products:
    add_bullet(b)

add_paragraph(
    'Стратегическая цель компании — стать ведущим разработчиком корпоративного '
    'программного обеспечения в Северо-Западном регионе России, последовательно '
    'расширяя клиентскую базу и портфель продуктов. Компания активно инвестирует '
    'в R&D, внедряя современные подходы к разработке: микросервисную архитектуру, '
    'DevOps-практики и облачные технологии.'
)

add_paragraph(
    'Практика проходила в отделе backend-разработки, который является ключевым '
    'техническим подразделением компании, обеспечивающим серверную инфраструктуру '
    'разрабатываемых платформ. Отдел состоит из команды разработчиков, '
    'DevOps-инженеров и аналитиков.'
)

add_paragraph('Отдел занимается следующими направлениями:')
dept_areas = [
    'Проектирование и разработка серверных приложений на NestJS / Node.js',
    'Проектирование реляционных баз данных PostgreSQL',
    'Разработка и документирование REST API',
    'Реализация систем аутентификации и авторизации',
    'Интеграция с внешними сервисами (CRM, платёжные системы, сервисы доставки)',
    'Контейнеризация и оркестрация сервисов (Docker, Docker Compose)',
    'Настройка CI/CD и мониторинга',
]
for b in dept_areas:
    add_bullet(b)

add_heading('1.2. Техническое и программное обеспечение', level=3)

add_paragraph(
    'В рамках выполнения задач практики компанией было предоставлено рабочее место, '
    'укомплектованное современным техническим и программным обеспечением. '
    'Рабочая станция оснащена процессором Intel Core i7-12700, 32 ГБ оперативной '
    'памяти и твердотельным накопителем объёмом 1 ТБ. Рабочее место подключено '
    'к корпоративной сети с выделенным интернет-каналом и VPN-доступом '
    'к внутренним ресурсам компании.'
)

add_paragraph(
    'Для разработки программного обеспечения использовалось следующее '
    'профессионально-ориентированное программное обеспечение:'
)

sw_items = [
    'Операционная система: Ubuntu 22.04 LTS',
    'Среда разработки (IDE): Visual Studio Code с расширениями ESLint, Prettier, GitLens, REST Client',
    'Языки программирования: TypeScript 5.x, JavaScript (ES2022), HTML5, CSS3',
    'Фреймворки бэкенда: Node.js 20 LTS, NestJS 10, TypeORM 0.3',
    'Фреймворки фронтенда: Next.js 14, React 18, Tailwind CSS 3',
    'СУБД: PostgreSQL 16',
    'Кэш: Redis 7',
    'Контейнеризация: Docker 24, Docker Compose v2',
    'Система управления версиями: Git, GitHub',
    'CI/CD: GitHub Actions',
    'Документирование API: Swagger / OpenAPI 3.0',
    'Средства обеспечения безопасности: JWT-аутентификация, bcrypt-хеширование паролей, HTTPS',
    'Инструменты тестирования: Jest, Supertest',
    'Менеджер пакетов: npm / pnpm',
]
for item in sw_items:
    add_bullet(item)

add_heading('1.3. Анализ существующих недостатков и проблем управления заказами', level=3)

add_paragraph(
    'В ходе анализа деятельности клиентов компании ООО «Парнас АйТ» и изучения '
    'типичных проблем в области управления заказами была выявлена совокупность '
    'системных недостатков, препятствующих эффективной автоматизации бизнес-процессов. '
    'Большинство организаций используют разрозненные инструменты — электронные таблицы, '
    'мессенджеры и разнородные локальные программы — что порождает следующие проблемы:'
)

problems = [
    'Отсутствие единого центра управления заказами — данные хранятся в разных системах, '
    'нет целостного представления о состоянии всех заказов в режиме реального времени.',
    'Ручное распределение заказов — менеджеры тратят значительное время на назначение '
    'исполнителей, что приводит к неравномерной нагрузке и просрочкам.',
    'Непрозрачность статусов — клиент не имеет возможности отслеживать ход выполнения '
    'заказа; коммуникация ведётся по телефону или электронной почте без фиксации договорённостей.',
    'Отсутствие истории изменений — нет журнала аудита, фиксирующего кто, когда и что '
    'изменил в заказе, что затрудняет разбор конфликтных ситуаций.',
    'Неэффективное уведомление участников — сотрудники узнают о назначении задачи '
    'с задержкой, что замедляет выполнение.',
    'Невозможность аналитики — отсутствие структурированных данных не позволяет строить '
    'отчёты по KPI сотрудников, среднему времени обработки и доходности.',
    'Сложность интеграции с внешними сервисами — нет API для подключения CRM, '
    'платёжных систем и служб доставки.',
]
for i, p in enumerate(problems, 1):
    add_numbered(p, i)

add_paragraph(
    'Выявленные недостатки явно отражают потребность в разработке специализированной '
    'веб-платформы OrderFlow, обеспечивающей полный цикл автоматизации обработки '
    'заказов: от момента поступления заявки до её выполнения, доставки и аналитики '
    'по результатам работы.'
)

add_heading('1.4. Формулировка цели и задач практики', level=3)

add_paragraph(
    'Целью практики является получение практического опыта в проектировании '
    'и разработке масштабируемых веб-платформ, применение теоретических знаний '
    'в области информационных систем, архитектуры ПО и современных технологий '
    'разработки, полученных в процессе обучения.'
)

add_paragraph('В ходе прохождения практики были решены следующие задачи:')

tasks = [
    'Провести анализ предметной области и сформулировать требования к системе управления заказами.',
    'Спроектировать реляционную схему базы данных PostgreSQL для хранения пользователей, '
    'заказов, клиентов, истории изменений и уведомлений.',
    'Разработать REST API на базе фреймворка NestJS с разграничением доступа по ролям '
    '(администратор, менеджер, исполнитель, клиент) с использованием JWT-аутентификации.',
    'Реализовать модуль автоматического распределения заказов с поддержкой нескольких '
    'алгоритмов: Round Robin, по минимальной загрузке, по приоритету.',
    'Разработать модуль уведомлений с поддержкой каналов Email, SMS и системных уведомлений.',
    'Создать интерактивный пользовательский интерфейс на Next.js / React с поддержкой '
    'полного цикла управления заказами, аналитических дашбордов и панели администратора.',
    'Настроить контейнеризацию приложения с помощью Docker Compose для воспроизводимого '
    'развёртывания в продакшн-среде.',
    'Настроить CI/CD-пайплайн на GitHub Actions для автоматизации сборки, тестирования '
    'и деплоя приложения.',
]
for i, t in enumerate(tasks, 1):
    add_numbered(t, i)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 2
# ═══════════════════════════════════════════════════════════════════════════
add_heading('Раздел 2', level=1)
add_heading(
    '2. СОДЕРЖАНИЕ МЕРОПРИЯТИЙ ПО РАЗРАБОТКЕ СИСТЕМЫ\nАВТОМАТИЗАЦИИ ЗАКАЗОВ',
    level=2
)
add_heading('2.1. Анализ и сравнительная оценка методов решения задач практики', level=3)

add_paragraph(
    'Для выбора оптимального технологического стека был проведён сравнительный анализ '
    'фреймворков серверной разработки на языке TypeScript / JavaScript. '
    'Рассматривались следующие варианты:'
)

add_paragraph(
    'Express.js — минималистичный и гибкий фреймворк, де-факто стандарт '
    'для Node.js-разработки. Обладает огромной экосистемой и широким сообществом, '
    'однако не имеет встроенной структуры проекта, что при масштабировании приводит '
    'к неоднородности кодовой базы. Отсутствие встроенного IoC-контейнера '
    'усложняет управление зависимостями в крупных проектах.',
    bold=False
)

add_paragraph(
    'Fastify — высокопроизводительный фреймворк с встроенной схемной валидацией '
    'через JSON Schema. Превосходит Express по скорости обработки запросов, '
    'однако имеет меньшее сообщество и более высокий порог вхождения. '
    'Не имеет встроенных инструментов для организации модульной архитектуры '
    'в стиле Angular/Spring.',
    bold=False
)

add_paragraph(
    'NestJS — прогрессивный фреймворк, построенный поверх Express (или Fastify), '
    'предоставляющий опinionated-архитектуру с поддержкой модульности, '
    'внедрения зависимостей, декораторов и встроенной поддержкой TypeScript. '
    'Имеет готовые решения для Guard, Interceptor, Pipe, Filter и тесную интеграцию '
    'с TypeORM, Swagger и другими библиотеками.',
    bold=False
)

add_paragraph(
    'В ходе сравнительного анализа был выбран NestJS, поскольку он обеспечивает '
    'структурированную модульную архитектуру, высокую поддерживаемость кода '
    'и нативную интеграцию со всеми необходимыми компонентами системы: TypeORM, '
    'Passport (JWT), Swagger, Redis, Bull (очереди задач).'
)

add_paragraph('Для фронтенд-разработки были рассмотрены следующие подходы:')

add_paragraph(
    'Create React App (CRA) — стандартный инструмент для создания React-приложений, '
    'однако в 2023–2024 годах признан устаревшим командой React. Медленная сборка, '
    'отсутствие SSR и неактуальная конфигурация webpack делают его '
    'нецелесообразным для новых проектов.',
    bold=False
)

add_paragraph(
    'Vue.js + Nuxt — удобный фреймворк с двусторонним связыванием данных. '
    'Nuxt предоставляет SSR из коробки, однако Vue имеет меньшую долю рынка '
    'и меньше готовых enterprise-UI компонентов по сравнению с React-экосистемой.',
    bold=False
)

add_paragraph(
    'Next.js + React — полнофункциональный фреймворк с поддержкой App Router, '
    'Server Components, SSR и SSG. Тесная интеграция с Vercel, отличная '
    'TypeScript-поддержка, мощная экосистема компонентов и инструментов. '
    'Наиболее зрелое решение для корпоративных веб-платформ.',
    bold=False
)

add_paragraph(
    'Для реализации фронтенда был выбран Next.js 14 совместно с React 18 '
    'и Tailwind CSS, как наиболее зрелое и высокопроизводительное решение, '
    'позволяющее эффективно строить сложные интерфейсы с управлением состоянием, '
    'серверным рендерингом и богатым дашбордом аналитики.'
)

add_paragraph(
    'Для выбора способа хранения данных проводился анализ СУБД. '
    'Рассматривались MySQL / MariaDB и PostgreSQL. '
    'PostgreSQL был выбран ввиду поддержки JSONB-полей для гибкого хранения '
    'метаданных заказов, мощного полнотекстового поиска, транзакционной надёжности '
    'и широкой поддержки в экосистеме TypeORM. '
    'Для кэширования результатов запросов и хранения сессий выбран Redis '
    'ввиду высокой производительности операций in-memory и встроенной поддержки '
    'в модуле @nestjs/cache-manager.'
)

add_heading('2.2. Описание реализации решения', level=3)

add_paragraph(
    'Программная реализация платформы OrderFlow разделена на три ключевых уровня:'
)

levels = [
    'Уровень данных — реляционная схема базы данных и ORM-модели',
    'Уровень бизнес-логики — REST API, модули автораспределения, уведомлений и аналитики',
    'Уровень представления — Next.js / React-приложение с компонентным интерфейсом',
]
for l in levels:
    add_bullet(l)

add_paragraph('1. Уровень данных. Схема базы данных.', bold=True, first_indent=False)

add_paragraph(
    'Реляционная схема базы данных включает следующие сущности, реализованные '
    'посредством TypeORM на основе декларативного стиля с декораторами:'
)

entities = [
    'users — пользователи системы с атрибутами: email, хеш пароля (bcrypt), '
    'ФИО, телефон, роль (enum: admin / manager / executor / client), '
    'флаг активности и дата регистрации.',
    'orders — заказы с полями: уникальный номер (auto-generated), '
    'статус (enum: new / accepted / processing / assigned / in_progress / ready / '
    'shipped / delivered / completed / cancelled), приоритет, сумма, '
    'срок выполнения, комментарий. Связи: many-to-one с customer, '
    'many-to-one с executor (User), one-to-many с order_items.',
    'order_items — позиции заказа: наименование товара / услуги, количество, '
    'цена, единица измерения.',
    'customers — карточка клиента: ФИО, компания, телефон, email, адрес, '
    'сегмент, примечания. Связь one-to-many с orders.',
    'order_history — журнал изменений заказов с фиксацией: кто изменил, '
    'когда, какое поле, старое и новое значение.',
    'notifications — уведомления пользователей: тип, заголовок, текст, '
    'флаг прочтения, канал доставки, дата.',
]
for e in entities:
    add_bullet(e)

add_paragraph(
    'Статусная модель заказа включает 10 состояний: '
    'new → accepted → processing → assigned → in_progress → ready → '
    'shipped → delivered → completed (или cancelled в любой точке). '
    'Допустимые переходы строго регламентированы словарём STATUS_TRANSITIONS '
    'в сервисном слое, что исключает некорректные изменения статуса как '
    'на уровне API, так и на уровне интерфейса.'
)

add_paragraph('2. Уровень бизнес-логики. REST API и модули системы.', bold=True, first_indent=False)

add_paragraph(
    'Серверная часть реализована на NestJS и организована в виде функциональных '
    'модулей, каждый из которых отвечает за свою область функциональности:'
)

modules = [
    'AuthModule — регистрация (POST /auth/register), вход (POST /auth/login), '
    'обновление токена (POST /auth/refresh). JWT-стратегия с passport-jwt. '
    'Роли защищены декоратором @Roles() и RolesGuard.',
    'OrdersModule — CRUD-операции с заказами: создание (POST /orders), '
    'список с фильтрацией (GET /orders), детали (GET /orders/:id), '
    'редактирование (PUT /orders/:id), смена статуса (PATCH /orders/:id/status), '
    'назначение исполнителя (PATCH /orders/:id/assign), история (GET /orders/:id/history).',
    'CustomersModule — управление клиентами: CRUD, поиск, фильтрация по сегменту, '
    'экспорт в CSV.',
    'UsersModule — управление пользователями: список, профиль, смена пароля, загрузка аватара.',
    'AutoAssignModule — сервис автоматического распределения заказов. '
    'Поддерживает три алгоритма: Round Robin (последовательное назначение), '
    'Min Load (исполнитель с минимальным числом активных заказов), '
    'Priority-based (учёт приоритета заказа и специализации исполнителя).',
    'NotificationsModule — отправка уведомлений по каналам Email (Nodemailer), '
    'системные in-app уведомления. Хранение уведомлений в БД, '
    'маркировка прочитанных.',
    'AnalyticsModule — аналитические отчёты: количество заказов по статусам, '
    'выручка по периодам, среднее время обработки, эффективность исполнителей, '
    'просроченные заказы, статистика клиентов.',
    'AdminModule — панель администратора: управление пользователями и ролями, '
    'блокировка аккаунтов, системные настройки.',
]
for m in modules:
    add_bullet(m)

add_paragraph(
    'Аутентификация реализована на базе JWT-токенов (access + refresh). '
    'При входе пользователь получает access-токен (срок действия 15 минут) '
    'и refresh-токен (7 дней). Access-токен передаётся в заголовке '
    'Authorization: Bearer <token> при каждом запросе. '
    'Хранение паролей осуществляется в виде bcrypt-хешей с коэффициентом сложности 12.'
)

add_paragraph(
    'Модуль автоматического распределения заказов является ключевой бизнес-функцией '
    'платформы. При создании нового заказа или смене статуса на "accepted" '
    'сервис AutoAssignService выбирает подходящего исполнителя в соответствии '
    'с настроенным в системе алгоритмом. Алгоритм Round Robin обеспечивает '
    'равномерную ротацию назначений. Алгоритм Min Load выбирает исполнителя '
    'с наименьшим числом активных задач (статусы: assigned, in_progress). '
    'Приоритетный алгоритм учитывает срочность заказа и специализацию исполнителя.'
)

add_paragraph('3. Уровень представления. Next.js / React-приложение.', bold=True, first_indent=False)

add_paragraph(
    'Фронтенд реализован в виде одностраничного Next.js-приложения (App Router) '
    'с серверными компонентами и включает следующие страницы и компоненты:'
)

pages = [
    'Страница входа и регистрации (AuthForm) — форма с валидацией через react-hook-form + zod, '
    'обработкой ошибок и отображением статуса запроса.',
    'Дашборд (dashboard/page.tsx) — главная страница с карточками KPI: '
    'количество заказов, выручка за период, средние сроки, незакрытые просрочки. '
    'Виджеты реализованы с помощью библиотеки Recharts.',
    'Список заказов (orders/page.tsx) — таблица всех заказов с фильтрацией '
    'по статусу, приоритету, клиенту, исполнителю и диапазону дат. '
    'Менеджер видит все заказы, исполнитель — только назначенные ему, '
    'клиент — только свои.',
    'Карточка заказа (orders/[id]/page.tsx) — детальная информация о заказе: '
    'позиции, статус, история изменений, кнопки смены статуса (в зависимости от роли), '
    'назначение исполнителя, прикреплённые файлы.',
    'Форма создания / редактирования заказа — мульти-шаговая форма '
    'с выбором клиента, добавлением позиций, установкой приоритета и срока.',
    'Управление клиентами (customers/page.tsx) — CRM-список с поиском, '
    'фильтрацией по сегменту и карточкой клиента с историей заказов.',
    'Аналитика (analytics/page.tsx) — интерактивные графики и диаграммы: '
    'динамика заказов (линейный график), распределение по статусам (круговая диаграмма), '
    'рейтинг исполнителей (столбчатая диаграмма), KPI-панели.',
    'Панель администратора (admin/page.tsx) — управление пользователями: '
    'создание, редактирование, блокировка, назначение ролей.',
    'Профиль пользователя (profile/page.tsx) — редактирование данных, '
    'смена пароля, настройки уведомлений.',
]
for pg in pages:
    add_bullet(pg)

add_paragraph(
    'Работа с сервером организована через централизованный axios-клиент (src/lib/api.ts) '
    'с перехватчиками (interceptors) для автоматической подстановки токена '
    'в заголовок Authorization и обработки 401-ошибок (автоматическое обновление токена '
    'через refresh-endpoint). Серверное состояние кэшируется с помощью TanStack Query '
    '(@tanstack/react-query), что исключает дублирование запросов и обеспечивает '
    'автоматическую инвалидацию кэша при мутациях данных.'
)

add_paragraph(
    'При совместном использовании всех трёх уровней была создана полнофункциональная '
    'веб-платформа OrderFlow, обеспечивающая сквозную автоматизацию обработки заказов — '
    'от момента создания заявки до завершения с фиксацией каждого этапа в системе, '
    'уведомлением участников и формированием аналитики по результатам работы.'
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════
add_heading('ЗАКЛЮЧЕНИЕ', level=1)

add_paragraph(
    'В ходе прохождения производственной практики в ООО «Парнас АйТ» '
    'были решены все поставленные задачи по проектированию и разработке '
    'веб-платформы OrderFlow — системы автоматизации полного цикла обработки '
    'заказов, обеспечивающей централизованное управление, прозрачный контроль '
    'статусов и аналитику по результатам бизнес-процессов.'
)

add_paragraph('В результате выполнения поставленных задач были достигнуты следующие результаты:')

results = [
    'Спроектирована реляционная база данных PostgreSQL, охватывающая все сущности '
    'предметной области: пользователи с ролевой моделью (admin / manager / executor / client), '
    'заказы с расширенным описанием и позициями, карточки клиентов, '
    'журнал истории изменений и таблица уведомлений. Схема реализована через TypeORM '
    'с поддержкой миграций для управления версиями структуры БД.',
    'Разработан REST API на базе NestJS с разграничением прав доступа по ролям. '
    'Реализована JWT-аутентификация с bcrypt-хешированием паролей, '
    'статусная модель заказов с контролем допустимых переходов, '
    'фильтрация и поиск заказов, история изменений, а также уведомления '
    'по каналам Email и in-app при смене статуса.',
    'Реализован модуль автоматического распределения заказов, поддерживающий '
    'три конфигурируемых алгоритма (Round Robin, Min Load, Priority), '
    'что позволяет администратору выбирать стратегию в зависимости '
    'от специфики бизнес-процессов.',
    'Разработан модуль аналитики с эндпоинтами для получения KPI-показателей: '
    'количество заказов, динамика выручки, средние сроки обработки, '
    'эффективность исполнителей и список просроченных заказов.',
    'Создан интерактивный пользовательский интерфейс на Next.js / React '
    'с поддержкой полного цикла управления заказами: от создания до завершения. '
    'Административная панель обеспечивает управление пользователями, '
    'аналитические дашборды включают интерактивные графики и диаграммы (Recharts).',
    'Настроена контейнеризация приложения с помощью Docker Compose '
    '(сервисы: PostgreSQL, Redis, backend, frontend) для воспроизводимого '
    'развёртывания в продакшн-среде.',
    'Настроен CI/CD-пайплайн на GitHub Actions с этапами: '
    'линтинг (ESLint), юнит-тестирование (Jest), сборка Docker-образов.',
]
for r in results:
    add_bullet(r)

add_paragraph(
    'Практическая значимость работы заключается в возможности непосредственного '
    'внедрения разработанной платформы OrderFlow в продуктовый портфель '
    'ООО «Парнас АйТ» в качестве типового коробочного решения для клиентов '
    'из сферы ритейла, логистики и производства. '
    'Разработанная микросервисно-ориентированная архитектура позволяет '
    'масштабировать систему горизонтально по мере роста числа пользователей '
    'и объёма обрабатываемых заказов. '
    'В перспективе платформа может быть расширена модулями: '
    'WebSocket-уведомления в реальном времени, интеграция с Telegram-ботом, '
    'мобильное приложение и интеграция с маркетплейсами.'
)

add_paragraph(
    'В процессе выполнения практической работы были закреплены и углублены '
    'профессиональные компетенции в области разработки корпоративных '
    'информационных систем: проектирование реляционных баз данных, '
    'построение RESTful API, реализация ролевого управления доступом, '
    'контейнеризация и настройка CI/CD, разработка современных '
    'React-интерфейсов с серверным рендерингом.'
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════════════
add_heading('СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ', level=1)

refs = [
    'Информационные технологии : учебник для вузов / В. В. Трофимов, О. П. Ильина, '
    'В. И. Кияев, Е. В. Трофимова ; отв. ред. В. В. Трофимов. – Москва : '
    'Издательство Юрайт, 2024. – 546 с. – (Высшее образование). – '
    'ISBN 978-5-534-02518-2.',
    'Тузовский, А. Ф. Объектно-ориентированное программирование : учебное пособие '
    'для вузов / А. Ф. Тузовский. – Москва : Издательство Юрайт, 2023. – 213 с. – '
    '(Высшее образование). – ISBN 978-5-534-16316-2.',
    'Тузовский, А. Ф. Проектирование и разработка web-приложений : учебное пособие '
    'для вузов / А. Ф. Тузовский. – Москва : Издательство Юрайт, 2023. – 218 с. – '
    '(Высшее образование). – ISBN 978-5-534-16300-1.',
    'NestJS Documentation [Электронный ресурс]. – URL: https://docs.nestjs.com/.',
    'TypeORM Documentation [Электронный ресурс]. – URL: https://typeorm.io/.',
    'Next.js Documentation [Электронный ресурс]. – URL: https://nextjs.org/docs.',
    'React Documentation [Электронный ресурс]. – URL: https://react.dev/.',
    'Docker Documentation [Электронный ресурс]. – URL: https://docs.docker.com/.',
    'GitHub Actions Documentation [Электронный ресурс]. – URL: https://docs.github.com/actions.',
    'ГОСТ Р 56939-2024 — Защита информации. Разработка безопасного программного '
    'обеспечения. Общие требования. Введён 20.12.2024.',
    'ГОСТ 34.602-2020 — Информационные технологии. Техническое задание на создание '
    'автоматизированной системы. Введён 01.01.2022.',
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent       = Cm(1.25)
    pf.first_line_indent = Cm(-1.25)
    pf.space_after       = Pt(3)
    run = p.add_run(f'{i}. {r}')
    set_font(run, size=12)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# APPENDIX A – DB entities
# ═══════════════════════════════════════════════════════════════════════════
add_heading('Приложение А', level=1)
add_paragraph(
    'Листинг 1 – Сущности базы данных (backend/src/entities/, фрагмент)',
    align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, first_indent=False
)

code_a = """// user.entity.ts
export enum UserRole {
  ADMIN    = 'admin',
  MANAGER  = 'manager',
  EXECUTOR = 'executor',
  CLIENT   = 'client',
}

@Entity('users')
export class User {
  @PrimaryGeneratedColumn()
  id: number;

  @Column({ unique: true })
  email: string;

  @Column()
  passwordHash: string;

  @Column()
  firstName: string;

  @Column()
  lastName: string;

  @Column({ nullable: true })
  phone: string;

  @Column({ type: 'enum', enum: UserRole, default: UserRole.CLIENT })
  role: UserRole;

  @Column({ default: true })
  isActive: boolean;

  @CreateDateColumn()
  createdAt: Date;
}

// order.entity.ts
export enum OrderStatus {
  NEW        = 'new',
  ACCEPTED   = 'accepted',
  PROCESSING = 'processing',
  ASSIGNED   = 'assigned',
  IN_PROGRESS= 'in_progress',
  READY      = 'ready',
  SHIPPED    = 'shipped',
  DELIVERED  = 'delivered',
  COMPLETED  = 'completed',
  CANCELLED  = 'cancelled',
}

export enum OrderPriority {
  LOW    = 'low',
  MEDIUM = 'medium',
  HIGH   = 'high',
  URGENT = 'urgent',
}

@Entity('orders')
export class Order {
  @PrimaryGeneratedColumn()
  id: number;

  @Column({ unique: true })
  orderNumber: string;

  @Column({ type: 'enum', enum: OrderStatus, default: OrderStatus.NEW })
  status: OrderStatus;

  @Column({ type: 'enum', enum: OrderPriority, default: OrderPriority.MEDIUM })
  priority: OrderPriority;

  @Column({ type: 'decimal', precision: 10, scale: 2, default: 0 })
  totalAmount: number;

  @Column({ nullable: true })
  comment: string;

  @Column({ nullable: true })
  deadline: Date;

  @ManyToOne(() => Customer, customer => customer.orders)
  customer: Customer;

  @ManyToOne(() => User, { nullable: true })
  executor: User;

  @OneToMany(() => OrderItem, item => item.order, { cascade: true })
  items: OrderItem[];

  @OneToMany(() => OrderHistory, h => h.order)
  history: OrderHistory[];

  @CreateDateColumn()
  createdAt: Date;

  @UpdateDateColumn()
  updatedAt: Date;
}"""

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
run = p.add_run(code_a)
run.font.name = 'Courier New'
run.font.size = Pt(9)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# APPENDIX B – Orders service (auto-assign)
# ═══════════════════════════════════════════════════════════════════════════
add_heading('Приложение Б', level=1)
add_paragraph(
    'Листинг 2 – Сервис автоматического распределения заказов '
    '(backend/src/modules/orders/auto-assign.service.ts, фрагмент)',
    align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, first_indent=False
)

code_b = """export type AssignAlgorithm = 'round-robin' | 'min-load' | 'priority';

@Injectable()
export class AutoAssignService {
  private rrIndex = 0;

  constructor(
    @InjectRepository(User)
    private usersRepo: Repository<User>,
    @InjectRepository(Order)
    private ordersRepo: Repository<Order>,
  ) {}

  async assign(order: Order, algorithm: AssignAlgorithm): Promise<User | null> {
    const executors = await this.usersRepo.find({
      where: { role: UserRole.EXECUTOR, isActive: true },
    });
    if (!executors.length) return null;

    switch (algorithm) {
      case 'round-robin':
        return this.roundRobin(executors);
      case 'min-load':
        return this.minLoad(executors);
      case 'priority':
        return this.priorityBased(executors, order);
    }
  }

  private roundRobin(executors: User[]): User {
    const user = executors[this.rrIndex % executors.length];
    this.rrIndex++;
    return user;
  }

  private async minLoad(executors: User[]): Promise<User> {
    const loads = await Promise.all(
      executors.map(async (u) => ({
        user: u,
        load: await this.ordersRepo.count({
          where: {
            executor: { id: u.id },
            status: In([OrderStatus.ASSIGNED, OrderStatus.IN_PROGRESS]),
          },
        }),
      })),
    );
    return loads.sort((a, b) => a.load - b.load)[0].user;
  }

  private async priorityBased(executors: User[], order: Order): Promise<User> {
    if (order.priority === OrderPriority.URGENT ||
        order.priority === OrderPriority.HIGH) {
      return this.minLoad(executors);
    }
    return this.roundRobin(executors);
  }
}"""

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
run = p.add_run(code_b)
run.font.name = 'Courier New'
run.font.size = Pt(9)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# APPENDIX C – Frontend orders page fragment
# ═══════════════════════════════════════════════════════════════════════════
add_heading('Приложение В', level=1)
add_paragraph(
    'Листинг 3 – Компонент списка заказов '
    '(frontend/src/app/(dashboard)/orders/page.tsx, фрагмент)',
    align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, first_indent=False
)

code_c = """'use client';

export default function OrdersPage() {
  const [filters, setFilters] = useState<OrderFilters>({});
  const { data, isLoading } = useOrders(filters);

  const statusColor: Record<OrderStatus, string> = {
    new:        'bg-gray-100 text-gray-700',
    accepted:   'bg-blue-100 text-blue-700',
    processing: 'bg-yellow-100 text-yellow-700',
    assigned:   'bg-purple-100 text-purple-700',
    in_progress:'bg-indigo-100 text-indigo-700',
    ready:      'bg-green-100 text-green-700',
    shipped:    'bg-teal-100 text-teal-700',
    delivered:  'bg-emerald-100 text-emerald-700',
    completed:  'bg-green-200 text-green-800',
    cancelled:  'bg-red-100 text-red-700',
  };

  return (
    <div className="p-6">
      <div className="flex items-center justify-between mb-6">
        <h1 className="text-2xl font-bold text-gray-900">Заказы</h1>
        <Link href="/orders/new">
          <Button>+ Новый заказ</Button>
        </Link>
      </div>

      <OrderFilters filters={filters} onChange={setFilters} />

      {isLoading ? (
        <div className="flex justify-center py-12">
          <Spinner />
        </div>
      ) : (
        <Table>
          <TableHeader>
            <TableRow>
              <TableHead>Номер</TableHead>
              <TableHead>Клиент</TableHead>
              <TableHead>Статус</TableHead>
              <TableHead>Приоритет</TableHead>
              <TableHead>Исполнитель</TableHead>
              <TableHead>Сумма</TableHead>
              <TableHead>Срок</TableHead>
              <TableHead></TableHead>
            </TableRow>
          </TableHeader>
          <TableBody>
            {data?.orders.map((order) => (
              <TableRow key={order.id}>
                <TableCell className="font-mono">{order.orderNumber}</TableCell>
                <TableCell>{order.customer?.firstName} {order.customer?.lastName}</TableCell>
                <TableCell>
                  <span className={\`px-2 py-1 rounded-full text-xs \${statusColor[order.status]}\`}>
                    {ORDER_STATUS_LABELS[order.status]}
                  </span>
                </TableCell>
                <TableCell>
                  <PriorityBadge priority={order.priority} />
                </TableCell>
                <TableCell>{order.executor?.firstName ?? '—'}</TableCell>
                <TableCell>{formatCurrency(order.totalAmount)}</TableCell>
                <TableCell>
                  <DeadlineCell deadline={order.deadline} />
                </TableCell>
                <TableCell>
                  <Link href={\`/orders/\${order.id}\`}>
                    <Button variant="ghost" size="sm">Открыть</Button>
                  </Link>
                </TableCell>
              </TableRow>
            ))}
          </TableBody>
        </Table>
      )}
    </div>
  );
}"""

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
run = p.add_run(code_c)
run.font.name = 'Courier New'
run.font.size = Pt(9)

# ─── Save ─────────────────────────────────────────────────────────────────
out = '/home/user/Test/Отчёт_OrderFlow_Парнас_АйТ.docx'
doc.save(out)
print(f'Saved: {out}')
